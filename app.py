"""
Student Translator MVP - Flask Application
A web app that translates student documents from English to Kinyarwanda, French, Swahili, or Arabic using OpenAI GPT-4o.
"""

import os
from flask import Flask, render_template, request, redirect, url_for, send_file, flash
from werkzeug.utils import secure_filename
from openai import OpenAI
import PyPDF2
from docx import Document

# Load environment variables from .env file if it exists
try:
    from dotenv import load_dotenv
    load_dotenv()
except ImportError:
    pass  # python-dotenv not installed, use environment variables directly

# Initialize Flask app
app = Flask(__name__)
app.secret_key = os.environ.get('SECRET_KEY', 'dev-secret-key-change-in-production')
app.config['MAX_CONTENT_LENGTH'] = 16 * 1024 * 1024  # 16MB max file size
app.config['UPLOAD_FOLDER'] = 'uploads'
app.config['TRANSLATIONS_FOLDER'] = 'translations'

# Allowed file extensions
ALLOWED_EXTENSIONS = {'pdf', 'docx'}

# Supported languages
LANGUAGES = {
    'kinyarwanda': 'Kinyarwanda',
    'french': 'French',
    'swahili': 'Swahili',
    'arabic': 'Arabic'
}

# Initialize OpenAI client
client = OpenAI(api_key=os.environ.get('OPENAI_API_KEY'))


def allowed_file(filename):
    """Check if the uploaded file has an allowed extension."""
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS


def extract_text_from_pdf(file_path):
    """Extract text content from a PDF file."""
    try:
        text = ""
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
        return text.strip()
    except Exception as e:
        raise Exception(f"Error extracting text from PDF: {str(e)}")


def extract_text_from_docx(file_path):
    """Extract text content from a DOCX file."""
    try:
        doc = Document(file_path)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        return text.strip()
    except Exception as e:
        raise Exception(f"Error extracting text from DOCX: {str(e)}")


def translate_text(text, target_language):
    """Translate text using OpenAI GPT-4o model."""
    if not text or not text.strip():
        raise Exception("No text to translate")
    
    # Create the translation prompt
    prompt = f"""You are a professional translator for educational documents.

Translate the following text from English to {target_language}.
- Keep the meaning exact and accurate.
- Use clear and natural language suitable for university students.
- Maintain any academic or technical terms as precisely as possible.
- Do not add explanations, only output the translated text.

Text to translate:
{text}"""
    
    try:
        response = client.chat.completions.create(
            model="gpt-4o",
            messages=[
                {"role": "system", "content": "You are a professional translator specializing in educational documents."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.3,
            max_tokens=4000
        )
        
        translated_text = response.choices[0].message.content.strip()
        return translated_text
    except Exception as e:
        raise Exception(f"Error during translation: {str(e)}")


def create_docx_file(text, output_path):
    """Create a DOCX file with the translated text."""
    try:
        doc = Document()
        
        # Add the translated text
        paragraphs = text.split('\n')
        for para_text in paragraphs:
            if para_text.strip():
                doc.add_paragraph(para_text)
        
        doc.save(output_path)
    except Exception as e:
        raise Exception(f"Error creating DOCX file: {str(e)}")


@app.route('/')
def index():
    """Render the main upload form page."""
    return render_template('index.html', languages=LANGUAGES)


@app.route('/translate', methods=['POST'])
def translate():
    """Handle file upload, translation, and return the translated document."""
    
    # Check if file was uploaded
    if 'file' not in request.files:
        flash('No file selected. Please upload a document.', 'error')
        return redirect(url_for('index'))
    
    file = request.files['file']
    target_language = request.form.get('language', '').lower()
    
    # Validate file
    if file.filename == '':
        flash('No file selected. Please choose a file to upload.', 'error')
        return redirect(url_for('index'))
    
    if not allowed_file(file.filename):
        flash('Invalid file type. Please upload a PDF or DOCX file.', 'error')
        return redirect(url_for('index'))
    
    # Validate language selection
    if target_language not in LANGUAGES:
        flash('Invalid language selected. Please choose Kinyarwanda, French, Swahili, or Arabic.', 'error')
        return redirect(url_for('index'))
    
    # Create uploads directory if it doesn't exist
    os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
    os.makedirs(app.config['TRANSLATIONS_FOLDER'], exist_ok=True)
    
    try:
        # Save uploaded file
        filename = secure_filename(file.filename)
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(file_path)
        
        # Extract text based on file type
        file_ext = filename.rsplit('.', 1)[1].lower()
        if file_ext == 'pdf':
            extracted_text = extract_text_from_pdf(file_path)
        elif file_ext == 'docx':
            extracted_text = extract_text_from_docx(file_path)
        else:
            raise Exception("Unsupported file type")
        
        # Check if text was extracted
        if not extracted_text or not extracted_text.strip():
            raise Exception("No text could be extracted from the document. The file may be empty or corrupted.")
        
        # Translate the text
        language_name = LANGUAGES[target_language]
        translated_text = translate_text(extracted_text, language_name)
        
        # Create output filename
        base_name = os.path.splitext(filename)[0]
        output_filename = f"{base_name}_translated.docx"
        output_path = os.path.join(app.config['TRANSLATIONS_FOLDER'], output_filename)
        
        # Create the translated DOCX file
        create_docx_file(translated_text, output_path)
        
        # Clean up uploaded file
        os.remove(file_path)
        
        # Redirect to success page with download link
        return redirect(url_for('success', filename=output_filename))
        
    except Exception as e:
        # Clean up on error
        if os.path.exists(file_path):
            os.remove(file_path)
        
        flash(f'Error processing document: {str(e)}', 'error')
        return redirect(url_for('index'))


@app.route('/success')
def success():
    """Render the success page with download link."""
    filename = request.args.get('filename', '')
    if not filename:
        flash('No file to download.', 'error')
        return redirect(url_for('index'))
    
    return render_template('success.html', filename=filename)


@app.route('/download/<filename>')
def download(filename):
    """Download the translated document."""
    try:
        file_path = os.path.join(app.config['TRANSLATIONS_FOLDER'], secure_filename(filename))
        if not os.path.exists(file_path):
            flash('File not found.', 'error')
            return redirect(url_for('index'))
        
        return send_file(file_path, as_attachment=True, download_name=filename)
    except Exception as e:
        flash(f'Error downloading file: {str(e)}', 'error')
        return redirect(url_for('index'))


if __name__ == '__main__':
    # Check if OpenAI API key is set
    if not os.environ.get('OPENAI_API_KEY'):
        print("WARNING: OPENAI_API_KEY environment variable not set!")
        print("Please set it before running the application.")
    
    # Run the Flask app
    app.run(debug=True, host='0.0.0.0', port=5000)

